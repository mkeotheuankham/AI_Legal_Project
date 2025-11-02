# ໄຟລ໌: backend/ingest.py (ສະບັບແກ້ໄຂ)

import os
import glob
import re
from tqdm import tqdm
import docx
from langchain_core.documents import Document

# [ ປັບປຸງ 1 ] - Import ຈາກ package ໃໝ່
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings # ປ່ຽນຈາກ SentenceTransformerEmbeddings

# --- Configuration (ປັບປຸງໃໝ່) ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, os.pardir))
SOURCE_DIRECTORY = os.path.join(PROJECT_ROOT, "source_documents")
PERSIST_DIRECTORY = os.path.join(SCRIPT_DIR, "db_vector")
EMBEDDING_MODEL = "intfloat/multilingual-e5-base" # ໃຊ້ Model ທີ່ໄວຂຶ້ນ

ARTICLE_REGEX = re.compile(r"^(ມາດຕາ\s*\d+.*)", re.IGNORECASE)

def load_and_chunk_documents() -> list[Document]:
    # ... (ເນື້ອໃນຟັງຊັນຄືເກົ່າ) ...
    all_chunks = []
    docx_files = glob.glob(os.path.join(SOURCE_DIRECTORY, "**", "*.docx"), recursive=True)
    
    if not docx_files:
        print(f"ບໍ່ພົບໄຟລ໌ .docx ໃນ: {SOURCE_DIRECTORY}")
        return []

    print(f"ກຳລັງປະມວນຜົນ {len(docx_files)} ໄຟລ໌ .docx...")

    for file_path in tqdm(docx_files, desc="Processing files"):
        try:
            doc = docx.Document(file_path)
            file_name = os.path.basename(file_path)
            
            current_article_title = "ບົດນຳ"
            current_article_content = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                match = ARTICLE_REGEX.match(text)
                
                if match:
                    if current_article_content:
                        content_str = "\n".join(current_article_content)
                        metadata = {
                            "source": file_name,
                            "article": current_article_title
                        }
                        all_chunks.append(Document(page_content=content_str, metadata=metadata))
                    
                    current_article_title = text
                    current_article_content = [text]
                
                else:
                    current_article_content.append(text)
            
            if current_article_content:
                content_str = "\n".join(current_article_content)
                metadata = {
                    "source": file_name,
                    "article": current_article_title
                }
                all_chunks.append(Document(page_content=content_str, metadata=metadata))

        except Exception as e:
            print(f"Error processing file {file_path}: {e}")
            
    return all_chunks

def main():
    print("ເລີ່ມຕົ້ນຂະບວນການ Ingestion...")
    
    # 1. ໂຫຼດ ແລະ ແບ່ງຂໍ້ມູນຕາມມາດຕາ
    documents = load_and_chunk_documents()
    
    if not documents:
        print("ບໍ່ພົບເອກະສານທີ່ຈະປະມວນຜົນ.")
        return

    print(f"ສ້າງ chunks ສຳເລັດ. ຈຳນວນ chunks ທັງໝົດ: {len(documents)}")

    # 2. ໂຫຼດ Embedding Model
    print(f"ກຳລັງໂຫຼດ Embedding Model: {EMBEDDING_MODEL}")
    
    # [ ປັບປຸງ 2 ] - ປ່ຽນ class ໃຫ້ເປັນ HuggingFaceEmbeddings
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={'device': 'cpu'}, # ບັງຄັບໃຊ້ CPU
        encode_kwargs={'normalize_embeddings': True}
    )

    # 3. ສ້າງ Vector Store (ChromaDB)
    print(f"ກຳລັງສ້າງ Vector Store ຢູ່ທີ່: {PERSIST_DIRECTORY}")
    
    # [ ປັບປຸງ 3 ] - ໃຊ້ Chroma (ທີ່ import ມາຈາກບ່ອນໃໝ່)
    vectordb = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=PERSIST_DIRECTORY
    )
    
    print("Persistent vector store ສຳເລັດ.")
    print("Ingestion complete!")

if __name__ == "__main__":
    main()